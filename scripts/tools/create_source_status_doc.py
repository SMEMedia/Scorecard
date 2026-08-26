from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[2]
OUTPUT_PATH = PROJECT_ROOT / "outputs" / "Monthly and Weekly Scorecard Automation Source Status.docx"
OLD_OUTPUT_PATH = PROJECT_ROOT / "outputs" / "Monthly Scorecard Automation Source Status.docx"


ROWS = [
    {
        "source": "GA4",
        "monthly": "Working. Populates SME Media Data B-J, Q-U, Z, and GA4 search clicks/impressions by searchTerm grouping; SME Media Data (Detail) B-K; SME Media Engagement Metrics B-K. June AM Return Users now comes from newVsReturning = returning. App metrics are correct and include iOS plus Android.",
        "weekly": "Working. Populates SME Media Data B-O, R, and GA4 search clicks/impressions by searchTerm grouping. Week ending 2026-07-04 AM Web Podcast Plays returns 78.",
        "status": "Green",
        "needed": "No current blocker for GA4 app metrics. Previous manual app counts did not account for Android. AM Web Podcast Plays needs Jake review before changing the method.",
    },
    {
        "source": "Google Search Console",
        "monthly": "Working. Populates Google Search Console search clicks/impressions, Ave Search Position, and Search CTR.",
        "weekly": "Working. Populates Google Search Console search clicks/impressions, Ave Search Position, and Search CTR. Latest weekly dry run reused saved snapshots.",
        "status": "Orange",
        "needed": "Keep GA4 and Search Console search count columns separate. Search position and CTR should continue to use Google Search Console.",
    },
    {
        "source": "YouTube",
        "monthly": "Working. Populates SME Media Data Y and AR; SME Media Engagement Metrics L-N. AB is formula-owned.",
        "weekly": "Working. Populates SME Media Data Q. T is formula-owned.",
        "status": "Green",
        "needed": "No current blocker. Subscriber values are snapshot-based; for exact month-end values, run promptly after close or refresh snapshots intentionally.",
    },
    {
        "source": "Libsyn",
        "monthly": "Working. Populates SME Media Data X using Libsyn Unique Downloads.",
        "weekly": "Working. Populates SME Media Data P using Libsyn Unique Downloads for Sunday-Friday, labeled by the Saturday Week Ending date.",
        "status": "Green",
        "needed": "No public API route found; current Playwright CSV export workaround is the working path. Keep Libsyn browser profile/cached CSV healthy.",
    },
    {
        "source": "App Store Connect",
        "monthly": "Working for iOS only. Filtered to Advanced Manufacturing app ID 6494275046 / SKU com.smemedia.ios. June maps 5 downloads, matching App Analytics.",
        "weekly": "Not applicable. Weekly scorecard has no app-download column.",
        "status": "Green",
        "needed": "Keep Advanced Manufacturing filters in place. App Analytics report-request API was probed, but current key cannot create analytics report requests.",
    },
    {
        "source": "Google Play Console",
        "monthly": "Working. Downloads the Play Console installs CSV for com.smemedia.android and maps June 2026 Android installs as 40.",
        "weekly": "Not applicable. Weekly scorecard has no app-download column.",
        "status": "Green",
        "needed": "No current blocker. Keep the service account reporting permission active and continue spot-checking Play Console export totals against the UI.",
    },
    {
        "source": "HubSpot",
        "monthly": "Working. Monthly folder/email logic authenticated and returns Manufacturing Weekly sent emails, excluding clones and out-of-month emails.",
        "weekly": "Implemented. Needs next weekly dry run after monthly validation to confirm current token scope applies to weekly report as well.",
        "status": "Green",
        "needed": "Continue validating totals against HubSpot. Current discrepancy is expected because the SBS email was not counted in previous manual totals, but should be included.",
    },
    {
        "source": "Meta: Facebook, Instagram, Threads",
        "monthly": "Scaffolded only. Target ownership: Facebook Followers AQ and Instagram Followers AU; Threads not mapped.",
        "weekly": "No weekly mapping implemented.",
        "status": "Red",
        "needed": "Need Meta developer access, app permissions, page/account IDs, and final metric mapping.",
    },
    {
        "source": "X",
        "monthly": "Scaffolded only. Target ownership: X Followers AT. API token exists but account has no credits.",
        "weekly": "No weekly mapping implemented.",
        "status": "Red",
        "needed": "Need paid X API credits or an alternate source/export for follower count.",
    },
    {
        "source": "LinkedIn",
        "monthly": "Planned only. Target ownership: SME Media Data (Detail) L-Y and LinkedIn Followers AS.",
        "weekly": "No weekly mapping implemented.",
        "status": "Red",
        "needed": "Need LinkedIn company verification: https://www.linkedin.com/developers/apps/verification/aa87d933-54de-4562-8543-dc3453404cb8",
    },
    {
        "source": "Walsworth Thermostats",
        "monthly": "Placeholder only. Target ownership: SME Media Data K-N.",
        "weekly": "Placeholder only.",
        "status": "Red",
        "needed": "Need credentials/API/export access and confirmation of exact fields.",
    },
    {
        "source": "Personify / Fonteva",
        "monthly": "Placeholder only. Target ownership: SME Media Data N-P.",
        "weekly": "Placeholder only.",
        "status": "Red",
        "needed": "Need access, API documentation, and field mapping.",
    },
    {
        "source": "DataBox",
        "monthly": "Not used for current automation. Previously explored for follower metrics.",
        "weekly": "Not used.",
        "status": "Orange",
        "needed": "Could not identify a reliable Databox API export path for existing metrics; native APIs are preferred.",
    },
    {
        "source": "IO / Pipeline",
        "monthly": "Not mapped yet.",
        "weekly": "Pending for SME Media Data AE-AO.",
        "status": "Red",
        "needed": "Need source system, access method, and field definitions for IO and pipeline values.",
    },
]


DISCREPANCY_ROWS = [
    {
        "priority": "1",
        "area": "Google Search Console",
        "difference": "Google Search Console Search Clicks +2,080 (+295%), Google Search Console Search Impressions +466,093 (+131%), Ave Search Position 16.88 vs 26.37.",
        "note": "Search count columns are now split between GA4 and Google Search Console. Search position and CTR remain Google Search Console metrics.",
    },
    {
        "priority": "2",
        "area": "GA4 app metrics",
        "difference": "App Sessions 121 vs 54, App Views 321 vs 103, App Users 54 vs 20, App New Users 21 vs 3, App Return Users 33 vs 17.",
        "note": "Automation metrics are correct because they include Android plus iOS; previous manual counts did not account for Android.",
    },
    {
        "priority": "3",
        "area": "HubSpot email metrics",
        "difference": "Email Opens +5,352 (+20.9%), Email Clicks +601 (+18.0%), Emails Delivered +8,950 (+2.6%).",
        "note": "SBS email was not counted in previous manual totals, but should be counted.",
    },
    {
        "priority": "4",
        "area": "AM Web Podcast Plays",
        "difference": "Automation 258 vs live 330, delta -72 (-21.8%).",
        "note": "GA4 exploration total 330 equals eventName = audio across all audio events. Current automation is narrower: eventName = audio plus pagePath contains /multimedia/podcasts/, returning 258. Counting only audio_player_action = play would be 98 broad or 78 with the podcast path. Consult Jake before changing the automation.",
    },
    {
        "priority": "5",
        "area": "Libsyn podcast downloads",
        "difference": "Previous automation used IAB Downloads and included Saturday. Live/manual uses Unique Downloads and a Sunday-Friday window for the Saturday Week Ending row.",
        "note": "Calculation changed to follow the live/manual method.",
    },
    {
        "priority": "6",
        "area": "GA4 AM web metrics",
        "difference": "AM Return Users now 2,276 vs live 2,251, delta +25 (+1.1%); other AM web deltas are under 1%.",
        "note": "Return-user implementation is now close after switching to newVsReturning = returning.",
    },
    {
        "priority": "7",
        "area": "Minor differences",
        "difference": "Monthly App Downloads +1, YouTube Podcast Plays -1, YouTube Subscribers -1, Email Subscribers -2.",
        "note": "Low-priority spot checks only.",
    },
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[float]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def style_paragraph(paragraph, size=8.2, color="111111", bold=False) -> None:
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor.from_string(color)
        run.bold = bold


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def status_fill(status: str) -> str:
    return {
        "Green": "D9EAD3",
        "Orange": "FCE5CD",
        "Red": "F4CCCC",
    }.get(status, "FFFFFF")


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.first_line_indent = Inches(-0.18)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run("- " + text)
    run.font.name = "Arial"
    run.font.size = Pt(9)


def build_doc() -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(1)
    title_run = title.add_run("Scorecard Automation Source Status")
    title_run.font.name = "Arial"
    title_run.font.size = Pt(21)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(31, 58, 95)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(7)
    today = date.today()
    run = subtitle.add_run(f"Monthly and Weekly automation snapshot as of {today.strftime('%B')} {today.day}, {today.year}")
    run.font.name = "Arial"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(90, 90, 90)

    summary = doc.add_paragraph()
    summary.paragraph_format.space_after = Pt(7)
    summary_run = summary.add_run(
        "Latest verification: Monthly dry run completed and planned one update for AM Web Podcast Plays; Weekly dry run completed with zero planned changes because saved snapshots already match the Google Sheet. Snapshot mode is active for saved writes."
    )
    summary_run.font.name = "Arial"
    summary_run.font.size = Pt(9)

    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["Source", "Monthly Status", "Weekly Status", "Stage", "What Is Needed"]
    widths = [1.25, 2.8, 2.55, 0.85, 2.45]

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, "1F3A5F")
        set_cell_margins(cell)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            style_paragraph(paragraph, size=8.4, color="FFFFFF", bold=True)

    for item in ROWS:
        cells = table.add_row().cells
        values = [item["source"], item["monthly"], item["weekly"], item["status"], item["needed"]]
        for idx, value in enumerate(values):
            cell = cells[idx]
            cell.text = value
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            set_cell_margins(cell)
            set_cell_shading(cell, status_fill(item["status"]) if idx == 3 else "FFFFFF")
            for paragraph in cell.paragraphs:
                if idx == 3:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    style_paragraph(paragraph, size=8.3, bold=True)
                else:
                    style_paragraph(paragraph)

    set_table_width(table, widths)

    notes_heading = doc.add_paragraph()
    notes_heading.paragraph_format.space_before = Pt(8)
    notes_heading.paragraph_format.space_after = Pt(3)
    h_run = notes_heading.add_run("Current Run Notes")
    h_run.font.name = "Arial"
    h_run.font.size = Pt(13)
    h_run.bold = True
    h_run.font.color.rgb = RGBColor(31, 58, 95)

    add_bullet(doc, "Snapshot mode stores period-specific records in config/state/scorecard_snapshots.json on saved runs. Dry runs do not create snapshots.")
    add_bullet(doc, "Monthly AM Web Podcast Plays currently maps GA4 eventCount where eventName = audio and pagePath contains /multimedia/podcasts/. June returned 258.")
    add_bullet(doc, "The GA4 exploration shown by Melanie returns 330 for June when counting eventName = audio across the Audio Events segment. The difference appears to be the page-path filter: eventName = audio alone returns 330, while adding pagePath contains /multimedia/podcasts/ returns 258. Counting only audio_player_action = play would return 98 without the page-path filter or 78 with it. Do not change the automation until Jake confirms the preferred calculation.")
    add_bullet(doc, "Weekly AM Web Podcast Plays uses the same GA4 definition. Week ending 2026-07-04 returned 78 and is already captured in snapshots.")
    add_bullet(doc, "App Store Connect is now filtered to Advanced Manufacturing only: app ID 6494275046 and SKU com.smemedia.ios. June maps to 5, matching App Analytics Total Downloads.")
    add_bullet(doc, "Google Play Console is now working. The latest monthly source run downloaded the Play Console installs CSV and mapped 40 Android installs for June 2026.")
    add_bullet(doc, "HubSpot monthly email metrics are now authenticating and returning records; continue validating mapped values against HubSpot.")
    add_bullet(doc, "Libsyn now uses Unique Downloads. Weekly Libsyn rows are labeled by the Saturday Week Ending date but sum Sunday through Friday to match the live/manual scorecard method.")

    discrepancy_heading = doc.add_paragraph()
    discrepancy_heading.paragraph_format.space_before = Pt(8)
    discrepancy_heading.paragraph_format.space_after = Pt(3)
    d_run = discrepancy_heading.add_run("Current Monthly Discrepancy Priorities")
    d_run.font.name = "Arial"
    d_run.font.size = Pt(13)
    d_run.bold = True
    d_run.font.color.rgb = RGBColor(31, 58, 95)

    discrepancy_note = doc.add_paragraph()
    discrepancy_note.paragraph_format.space_after = Pt(4)
    dn_run = discrepancy_note.add_run(
        "Read-only comparison of current automated Monthly source values against the live Monthly Scorecard row for June 2026."
    )
    dn_run.font.name = "Arial"
    dn_run.font.size = Pt(9)

    discrepancy_table = doc.add_table(rows=1, cols=4)
    discrepancy_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    discrepancy_table.style = "Table Grid"
    discrepancy_headers = ["Priority", "Area", "Difference", "Current Note / Holdup"]
    discrepancy_widths = [0.65, 1.7, 4.15, 3.4]

    for idx, header in enumerate(discrepancy_headers):
        cell = discrepancy_table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, "1F3A5F")
        set_cell_margins(cell)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            style_paragraph(paragraph, size=8.4, color="FFFFFF", bold=True)

    for item in DISCREPANCY_ROWS:
        cells = discrepancy_table.add_row().cells
        values = [item["priority"], item["area"], item["difference"], item["note"]]
        for idx, value in enumerate(values):
            cell = cells[idx]
            cell.text = value
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            set_cell_margins(cell)
            if idx == 0:
                set_cell_shading(cell, "FCE5CD")
            else:
                set_cell_shading(cell, "FFFFFF")
            for paragraph in cell.paragraphs:
                if idx == 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    style_paragraph(paragraph, size=8.3, bold=True)
                else:
                    style_paragraph(paragraph)

    set_table_width(discrepancy_table, discrepancy_widths)

    link_heading = doc.add_paragraph()
    link_heading.paragraph_format.space_before = Pt(6)
    link_heading.paragraph_format.space_after = Pt(2)
    l_run = link_heading.add_run("Key External Link")
    l_run.font.name = "Arial"
    l_run.font.size = Pt(12)
    l_run.bold = True
    l_run.font.color.rgb = RGBColor(31, 58, 95)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.add_run("LinkedIn company verification: ")
    add_hyperlink(
        p,
        "linkedin.com/developers/apps/verification/aa87d933-54de-4562-8543-dc3453404cb8",
        "https://www.linkedin.com/developers/apps/verification/aa87d933-54de-4562-8543-dc3453404cb8",
    )
    for run in p.runs:
        run.font.name = "Arial"
        run.font.size = Pt(9)

    doc.save(OUTPUT_PATH)
    if OLD_OUTPUT_PATH.exists():
        OLD_OUTPUT_PATH.unlink()


if __name__ == "__main__":
    build_doc()
    print(OUTPUT_PATH)
